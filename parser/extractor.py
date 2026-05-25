"""
extractor.py
Core extraction pipeline — NER (spaCy) + regex hybrid layer.
Phase 6: adds per-entity confidence scores derived from eval F1 metrics.
"""

import json
import os
import re

import spacy

from .regex_patterns import (
    EMAIL,
    PHONE,
    LINKEDIN,
    GITHUB,
    LINKEDIN_USERNAME,
    GITHUB_USERNAME,
    GRADUATION_YEAR,
    SKILLS_KEYWORDS,
    DEGREE_KEYWORDS,
)


class ResumeExtractor:
    """Full resume parsing pipeline combining spaCy NER and regex patterns."""

    def __init__(self, model_path: str = "training/output/model-best"):
        """Load the custom spaCy model from *model_path*.

        Also loads eval_results.json once so per-label F1 scores can be used
        as pseudo-confidence proxies for NER predictions.
        """
        self.nlp = spacy.load(model_path)
        self.label_f1 = self._load_label_f1(model_path)

    # ─── Eval F1 loading ─────────────────────────────────────────────────────

    @staticmethod
    def _load_label_f1(model_path: str) -> dict:
        """Return a mapping {label: f1_score_0_to_100} from eval_results.json.

        Searches upward from model_path for the file.  Returns an empty dict
        if the file cannot be found or parsed.
        """
        # Walk up directories from model_path to find training/eval_results.json
        candidates = []
        path = os.path.abspath(model_path)
        for _ in range(5):
            path = os.path.dirname(path)
            candidate = os.path.join(path, "training", "eval_results.json")
            candidates.append(candidate)
            candidate2 = os.path.join(path, "eval_results.json")
            candidates.append(candidate2)

        for candidate in candidates:
            if os.path.exists(candidate):
                try:
                    with open(candidate, encoding="utf-8") as f:
                        data = json.load(f)
                    per_label = data.get("per_label", {})
                    return {lbl: v["f1"] * 100.0 for lbl, v in per_label.items()}
                except Exception:
                    pass
        return {}

    # ─── Text Extraction ──────────────────────────────────────────────────────

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract plain text from a PDF using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
        except ImportError as exc:
            raise ImportError(
                "PyMuPDF is required to parse PDFs. "
                "Install it with: pip install PyMuPDF"
            ) from exc

        try:
            doc = fitz.open(file_path)
        except Exception as exc:
            raise ValueError(
                f"Could not open PDF file '{file_path}': {exc}"
            ) from exc

        pages = []
        for page in doc:
            page_text = page.get_text("text")

            # PASS 2: Links
            for link in page.get_links():
                if link.get("kind") in (2, getattr(fitz, "LINK_URI", 2)):
                    uri = link.get("uri", "")
                    if uri:
                        page_text += f"\n{uri}"

            # PASS 3: Annotations
            try:
                for annot in page.annots():
                    if annot.type[0] == fitz.PDF_ANNOT_LINK:
                        action = annot.info.get("action", "")
                        if action.startswith("http") or "linkedin" in action or "github" in action:
                            page_text += f"\n{action}"
            except Exception:
                pass

            # PASS 4: Raw Dict Extraction
            page_dict = page.get_text("dict")
            for block in page_dict.get("blocks", []):
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        uri = span.get("uri", "")
                        if uri:
                            page_text += f"\n{uri}"

            pages.append(page_text)
        doc.close()

        return "\n\n".join(pages).strip()

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract plain text from a .docx file using python-docx."""
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError(
                "python-docx is required to parse .docx files. "
                "Install it with: pip install python-docx"
            ) from exc

        doc = Document(file_path)

        # Paragraphs
        para_text = "\n".join(p.text for p in doc.paragraphs)

        # Tables
        table_parts = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        table_parts.append(cell.text.strip())
        table_text = "\n".join(table_parts)

        combined = "\n".join(filter(None, [para_text, table_text]))
        return combined.strip()

    def extract_text(self, file_path: str) -> str:
        """Route text extraction to the correct method based on file extension."""
        ext = os.path.splitext(file_path)[-1].lower()

        if ext == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return self.extract_text_from_docx(file_path)
        elif ext == ".txt":
            with open(file_path, encoding="utf-8") as f:
                return f.read().strip()
        else:
            raise ValueError(f"Unsupported file format: '{ext}'")

    # ─── NER ─────────────────────────────────────────────────────────────────

    def _run_ner(self, text: str) -> dict:
        """Run the custom spaCy model and collect entities by label.

        Each entity is stored as a dict::

            {"text": "entity text", "score": 0.9195}

        The score is derived from the per-label F1 loaded from
        eval_results.json — a meaningful proxy: high-F1 labels
        (Name ≈ 0.92) get high scores, low-F1 labels (Skills ≈ 0.20) get low.
        """
        result = {
            "name": [],
            "designation": [],
            "companies": [],
            "college_name": [],
            "degree": [],
            "graduation_year": [],
            "skills": [],
            "email": [],
            "location": [],
        }

        label_map = {
            "Name": "name",
            "Designation": "designation",
            "Companies worked at": "companies",
            "College Name": "college_name",
            "Degree": "degree",
            "Graduation Year": "graduation_year",
            "Skills": "skills",
            "Email Address": "email",
            "Location": "location",
        }

        doc = self.nlp(text)
        for ent in doc.ents:
            key = label_map.get(ent.label_)
            if key is not None:
                value = ent.text.strip()
                if value:
                    # Use label F1 as pseudo-confidence (0–100 → 0.0–1.0)
                    raw_f1 = self.label_f1.get(ent.label_, 50.0)
                    score = round(raw_f1 / 100.0, 4)
                    result[key].append({"text": value, "score": score})

        return result

    # ─── Regex ────────────────────────────────────────────────────────────────

    def _run_regex(self, text: str) -> dict:
        """Run all regex patterns on *text* and return structured results.

        Regex matches are deterministic, so they receive score = 1.0.
        Each item is a dict: {"text": ..., "score": 1.0}
        """
        result = {
            "name": [],
            "designation": [],
            "companies": [],
            "college_name": [],
            "degree": [],
            "graduation_year": [],
            "skills": [],
            "email": [],
            "location": [],
            # Regex-only fields
            "phone": [],
            "linkedin": [],
            "github": [],
        }

        def _item(text_val: str) -> dict:
            return {"text": text_val, "score": 1.0}

        # Email — filter out tracker URLs
        for match in EMAIL.findall(text):
            if "indeed.com" not in match.lower() and "linkedin.com" not in match.lower():
                result["email"].append(_item(match))

        # Phone — reconstruct full match from groups
        for groups in PHONE.findall(text):
            full = "".join(groups).strip()
            if full:
                result["phone"].append(_item(full))

        # LinkedIn
        for match in LINKEDIN.finditer(text):
            result["linkedin"].append(_item(match.group(0)))
            
        if not result["linkedin"]:
            for match in LINKEDIN_USERNAME.finditer(text):
                username = match.group(1)
                if username.lower() not in ["com", "www", "http", "https", "in", "jobs", "company", "pub"]:
                    result["linkedin"].append(_item("https://linkedin.com/in/" + username))

        # GitHub
        for match in GITHUB.finditer(text):
            result["github"].append(_item(match.group(0)))
            
        if not result["github"]:
            for match in GITHUB_USERNAME.finditer(text):
                username = match.group(1)
                if username.lower() not in ["com", "www", "http", "https", "in", "jobs", "company", "pub"]:
                    result["github"].append(_item("https://github.com/" + username))

        # Graduation year
        for year in GRADUATION_YEAR.findall(text):
            result["graduation_year"].append(_item(year))

        # Skills — case-insensitive keyword search
        text_lower = text.lower()
        for skill in SKILLS_KEYWORDS:
            if re.search(r'\b' + re.escape(skill.lower()) + r'\b', text_lower):
                result["skills"].append(_item(skill))

        # Degree keywords
        for degree in DEGREE_KEYWORDS:
            if degree.lower() in text_lower:
                result["degree"].append(_item(degree))

        return result

    # ─── Merge ────────────────────────────────────────────────────────────────

    def _merge_results(self, ner: dict, regex: dict) -> dict:
        """Merge NER and regex results with field-specific priority rules.

        Items are now dicts {"text": ..., "score": ...}.
        Deduplication compares on the "text" field (case-insensitive) and
        keeps the entry with the higher score.
        """

        def dedup_scored(lst: list) -> list:
            """Deduplicate by text (case-insensitive), keeping highest score."""
            best: dict = {}  # text_lower → item dict
            for item in lst:
                key = item["text"].lower()
                if key not in best or item["score"] > best[key]["score"]:
                    best[key] = item
            return list(best.values())

        def dedup_scored_ci(lst: list) -> list:
            """Same as dedup_scored — alias for readability."""
            return dedup_scored(lst)

        merged = {}

        # Skills: combine regex (score=1.0) + NER, dedup, sort alphabetically
        combined_skills = regex["skills"] + ner["skills"]
        merged["skills"] = sorted(
            dedup_scored(combined_skills),
            key=lambda x: x["text"].lower(),
        )

        # Graduation year: combine, dedup, keep valid 4-digit years 1980–2030
        combined_years = dedup_scored(ner["graduation_year"] + regex["graduation_year"])
        merged["graduation_year"] = [
            y for y in combined_years
            if y["text"].isdigit() and 1980 <= int(y["text"]) <= 2030
        ]

        # Email: combine, dedup case-insensitively
        merged["email"] = dedup_scored_ci(ner["email"] + regex["email"])

        # Degree: combine, dedup
        merged["degree"] = dedup_scored(ner["degree"] + regex["degree"])

        # NER-primary fields: fall back to regex if NER is empty
        for field in ("name", "designation", "companies", "college_name", "location"):
            ner_vals = dedup_scored(ner.get(field, []))
            if ner_vals:
                merged[field] = ner_vals
            else:
                merged[field] = dedup_scored(regex.get(field, []))

        # Name: keep only the first (highest-confidence) value
        if len(merged["name"]) > 1:
            merged["name"] = [merged["name"][0]]

        # Regex-only fields
        merged["phone"] = dedup_scored(regex.get("phone", []))
        merged["linkedin"] = dedup_scored(regex.get("linkedin", []))
        merged["github"] = dedup_scored(regex.get("github", []))

        return merged

    # ─── Public API ───────────────────────────────────────────────────────────

    def parse(self, file_path: str) -> dict:
        """Full pipeline: extract text → NER → regex → merge → return dict.

        All entity fields are lists of dicts: {"text": ..., "score": ...}
        """
        text = self.extract_text(file_path)
        ner_results = self._run_ner(text)
        regex_results = self._run_regex(text)
        merged = self._merge_results(ner_results, regex_results)

        merged["file_name"] = os.path.basename(file_path)
        merged["file_type"] = os.path.splitext(file_path)[-1].lstrip(".")
        merged["raw_text_length"] = len(text)

        return merged

    def parse_from_text(self, text: str) -> dict:
        """Parse raw text directly (e.g. when pasted in Streamlit UI).

        All entity fields are lists of dicts: {"text": ..., "score": ...}
        """
        ner_results = self._run_ner(text)
        regex_results = self._run_regex(text)
        merged = self._merge_results(ner_results, regex_results)

        merged["file_name"] = "pasted_text"
        merged["file_type"] = "text"
        merged["raw_text_length"] = len(text)

        return merged

    # ─── Utility ──────────────────────────────────────────────────────────────

    @staticmethod
    def get_flat_result(parsed: dict) -> dict:
        """Convert scored dict format back to simple flat lists of strings.

        Input:  {"name": [{"text": "John", "score": 0.92}], ...}
        Output: {"name": ["John"], ...}

        Non-list fields (file_name, file_type, raw_text_length) are passed
        through unchanged.
        """
        flat = {}
        skip_keys = {"file_name", "file_type", "raw_text_length"}
        for key, value in parsed.items():
            if key in skip_keys:
                flat[key] = value
            elif isinstance(value, list):
                flat[key] = [
                    item["text"] if isinstance(item, dict) else item
                    for item in value
                ]
            else:
                flat[key] = value
        return flat
